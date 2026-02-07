#!/usr/bin/env python3
"""Convert LaTeX conference paper to Word document format."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def clean_latex(text):
    """Remove LaTeX commands and clean text."""
    # Clean up LaTeX escaped characters FIRST (before removing comments)
    text = text.replace('\\%', '%')
    text = text.replace('\\$', '$')
    text = text.replace('\\&', '&')
    text = text.replace('\\_', '_')
    text = text.replace('\\#', '#')
    
    # Remove comments after unescaping %
    text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)
    
    # Remove section commands but keep section names
    text = re.sub(r'\\section\*?\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\subsection\*?\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\subsubsection\*?\{([^}]+)\}', r'\1', text)
    
    # Remove figure/table environments but keep captions
    text = re.sub(r'\\begin\{figure\}.*?\\caption\{([^}]+)\}.*?\\end\{figure\}', r'[Figure: \1]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{table\}.*?\\caption\{([^}]+)\}.*?\\end\{table\}', r'[Table: \1]', text, flags=re.DOTALL)
    
    # Remove enumerate/itemize environments
    text = re.sub(r'\\begin\{enumerate\}', '', text)
    text = re.sub(r'\\end\{enumerate\}', '', text)
    text = re.sub(r'\\begin\{itemize\}', '', text)
    text = re.sub(r'\\end\{itemize\}', '', text)
    text = re.sub(r'\\item\s*', '• ', text)
    
    # Handle custom commands (replace with meaningful text)
    text = re.sub(r'\\thillmo\{\}', 'thiLLMo', text)
    text = re.sub(r'\\ograg\{\}', 'OG-RAG', text)
    
    # Handle text formatting commands - extract content
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]+)\}', r'\1', text)
    
    # Handle citations and references
    text = re.sub(r'\\cite\{[^}]+\}', '[citation]', text)
    text = re.sub(r'\\ref\{[^}]+\}', '[reference]', text)
    text = re.sub(r'\\label\{[^}]+\}', '', text)
    
    # Handle URLs
    text = re.sub(r'\\url\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\href\{([^}]+)\}\{([^}]+)\}', r'\2 (\1)', text)
    
    # Handle footnotes
    text = re.sub(r'\\footnote\{([^}]+)\}', r' [\1]', text)
    
    # Remove math mode delimiters but keep content
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    
    # Remove remaining commands with braces - extract content only
    text = re.sub(r'\\[a-zA-Z@]+\*?\{([^}]*)\}', r'\1', text)
    
    # Remove standalone commands (no braces)
    text = re.sub(r'\\[a-zA-Z@]+\*?\s*', '', text)
    
    # Clean up braces
    text = text.replace('{', '')
    text = text.replace('}', '')
    
    # Clean up special characters
    text = text.replace('~', ' ')
    text = text.replace('``', '"')
    text = text.replace("''", '"')
    text = text.replace('---', '—')
    text = text.replace('--', '–')
    
    # Remove any remaining backslashes (catch-all)
    text = text.replace('\\', '')
    
    # Remove multiple spaces and blank lines
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n\n+', '\n\n', text)
    text = re.sub(r'\n ', '\n', text)
    
    return text.strip()

def extract_section(content, start_marker, end_marker=None):
    """Extract section between markers."""
    start_idx = content.find(start_marker)
    if start_idx == -1:
        return ""
    
    start_idx += len(start_marker)
    
    if end_marker:
        end_idx = content.find(end_marker, start_idx)
        if end_idx == -1:
            return content[start_idx:].strip()
        return content[start_idx:end_idx].strip()
    
    return content[start_idx:].strip()

def read_section_file(filepath):
    """Read and clean a section file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        # Remove begin/end document markers if present
        content = re.sub(r'\\begin\{document\}', '', content)
        content = re.sub(r'\\end\{document\}', '', content)
        return clean_latex(content)
    except FileNotFoundError:
        return ""

def create_word_document():
    """Create Word document from LaTeX sources."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('thiLLMo: Ontology-Grounded RAG for Culturally Faithful Kikuyu-English Proverb Translation')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # Abstract
    abstract_heading = doc.add_heading('Abstract', level=1)
    abstract_text = read_section_file('sections/00-abstract.tex')
    if abstract_text:
        doc.add_paragraph(abstract_text)
    
    # Introduction
    doc.add_heading('1. Introduction', level=1)
    intro_text = read_section_file('sections/01-introduction.tex')
    if intro_text:
        # Split by subsections if present
        paragraphs = intro_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
    
    # Related Work
    doc.add_heading('2. Related Work', level=1)
    related_text = read_section_file('sections/02-related-work.tex')
    if related_text:
        paragraphs = related_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
    
    # Methodology
    doc.add_heading('3. Methodology', level=1)
    method_text = read_section_file('sections/03-methodology.tex')
    if method_text:
        paragraphs = method_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
    
    # Results
    doc.add_heading('4. Results', level=1)
    results_text = read_section_file('sections/04-results.tex')
    if results_text:
        paragraphs = results_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
    
    # Discussion
    doc.add_heading('5. Discussion', level=1)
    discussion_text = read_section_file('sections/05-discussion.tex')
    if discussion_text:
        paragraphs = discussion_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
    
    # Conclusion
    doc.add_heading('6. Conclusion', level=1)
    conclusion_text = read_section_file('sections/06-conclusion.tex')
    if conclusion_text:
        paragraphs = conclusion_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
    
    # Add note about references
    doc.add_page_break()
    doc.add_heading('References', level=1)
    doc.add_paragraph('[References from references.bib - see PDF version for complete bibliography]')
    
    # Save document
    output_file = 'thiLLMo_DoCEIS2026_Conference_Paper.docx'
    doc.save(output_file)
    print(f"✓ Word document created: {output_file}")
    return output_file

if __name__ == '__main__':
    create_word_document()
