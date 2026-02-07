#!/usr/bin/env python3
"""Check Word document for LaTeX residue."""

from docx import Document

doc = Document('thiLLMo_DoCEIS2026_Conference_Paper.docx')
print('=== CHECKING WORD DOCUMENT FOR LATEX COMMANDS ===\n')

issues_found = 0
for i, para in enumerate(doc.paragraphs[:20]):
    text = para.text.strip()
    if text:
        # Check for LaTeX commands
        if '\\' in text:
            print(f'⚠️  Paragraph {i+1}: Found backslash')
            print(f'   {text[:200]}...\n')
            issues_found += 1
        elif '{' in text or '}' in text:
            print(f'⚠️  Paragraph {i+1}: Found braces')
            print(f'   {text[:200]}...\n')
            issues_found += 1

if issues_found == 0:
    print('✅ No LaTeX residue found in first 20 paragraphs!')
    print('\n=== SAMPLE CONTENT ===')
    for i, para in enumerate(doc.paragraphs[2:7]):
        if para.text.strip():
            print(f'\nParagraph {i+3}:')
            print(para.text[:250])
else:
    print(f'\n❌ Found {issues_found} paragraphs with LaTeX residue')
